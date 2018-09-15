from docx import Document
from docx.shared import Pt,RGBColor,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
from db import Db

class Db2Doc(object):

    def __init__(self):
        self.document=Document()
        self.document.add_heading('数据库设计文档', 0)
        gen_date=datetime.datetime.now().strftime('%Y-%m-%d')
        self.document.add_paragraph(gen_date)
        self.document.styles['Normal'].font.name = u'宋体'

        self.except_tables=[]

    def writeTitle(self,name):
        self.document.add_paragraph('')
        self.document.add_heading(name, 2)

    def writeTable(self,rows):
        rowCount=len(rows)+1
        colsCount=6
        table = self.document.add_table(rows=rowCount, cols=colsCount, style="Table Grid")
        table.autofit = False

        arr = [u'序号', u"字段名", u"类型", u"是否为空", u"主键", u"说明"]
        heading_cells = table.rows[0].cells
        for i in range(colsCount):
            p = heading_cells[i].paragraphs[0]
            run = p.add_run(arr[i])
            run.font.size = Pt(12)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for index in range(rowCount-1):
            table.cell(index+1, 0).text = str(index+1)
            p=table.cell(index+1, 0).paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row=rows[index]
            table.cell(index + 1, 1).text = row[0]
            table.cell(index + 1, 2).text = row[1]
            table.cell(index + 1, 3).text = row[3]
            table.cell(index + 1, 4).text = row[4]
            table.cell(index + 1, 5).text = row[8]

    def except_views(self,tbs):
        tbs=filter(lambda x:x[1]!="VIEW",tbs)
        return tbs

    def except_demos(self,tbs):
        tbs = filter(lambda x: x[0].startswith("demo"), tbs)
        return tbs

    def gen_doc(self):
        db = Db()
        tbs = db.getTables()

        self.except_demos(tbs)
        self.except_views(tbs)

        tbs=list(tbs)
        tbs.sort(reverse=True)

        i=0
        for tb in tbs:
            i=i+1
            self.writeTitle('%s.%s(%s)' %(str(i),tb[0],tb[1]))
            fds = db.getFullFields(tb[0])
            self.writeTable(fds)
        db.close()

        self.document.save('../output/数据库设计文档.docx')
        print('======>文档生成完成.')


doc=Db2Doc()

doc.gen_doc()

